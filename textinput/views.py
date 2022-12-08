from __future__ import (
    absolute_import, division, print_function, unicode_literals
)
from dataclasses import replace
from django.shortcuts import render, HttpResponseRedirect
from django.http import HttpResponse
from django.conf import settings
from django.core.files.storage import FileSystemStorage
from django.core.files import File
import os.path
import os
#from gtts import gTTS
#from playsound import  playsound
import threading
import re
import nltk
from nltk.corpus import stopwords 
from nltk.tokenize import word_tokenize, sent_tokenize
import googletrans
from googletrans import Translator
from time import sleep
import time
#import PyPDF2  
#import docx2txt
import io
import csv
import pandas as pd
#from fpdf import FPDF
import string
import docx
#import xlrd 
import openpyxl
#from google.cloud import translate
#import xlwt
#from joblib import Parallel, delayed
import time
from django.http import StreamingHttpResponse
from wsgiref.util import FileWrapper
#import pdfcrowd
import sys
import mimetypes
#import convertapi
#from pdf2docx import Converter
from django.contrib import messages
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
#from pdf2docx import parse
import mammoth
import pypandoc
from bs4 import BeautifulSoup, Tag,NavigableString
from modules.h2d import HtmlToDocx
#from pdf2docx import Converter
#from gtts import gTTS
#from playsound import  playsound

# Create your views here.
#from modules import googletrans
#from modules.googletrans import Translator
#translate_client = translate.Client()




#Read an excel file



# Function to extract the proper nouns 
def ProperNounExtractor(text):
    
    #print('PROPER NOUNS EXTRACTED :')
    
    #sentences = nltk.sent_tokenize(text)
    #for sentence in sentences:
        
        #words = nltk.word_tokenize(sentence)
        #words = [word for word in words if word not in set(stopwords.words('english'))]
    tagged = nltk.pos_tag(["Ribbons and Balloons is a cake shop"])
    tag=tagged[0][1]
    if(tag=='NNP'):
        print(text)
        return True
    '''for (text, tag) in tagged:
        if tag == 'NNP': # If the word is a proper noun
            print(text)
            return 
'''
'''text =  "Rohan is a wonderful player. He was born in India. He is a fan of the movie Wolverine. He has a dog named Bruno."
'''
# Calling the ProperNounExtractor function to extract all the proper nouns from the given text. 


x = "temp"
def home(request):
    global x
    x = "temp"
    #print(googletrans.LANGUAGES)
    lang =googletrans.LANGUAGES
    translator = Translator()
    truecheck = False 
    if request.method == 'POST':
        myfile = request.FILES['myfile']
        x=myfile.name
        src1 = request.POST['src']
        dest1 = request.POST['dest']
        #intype=request.POST
        names = myfile.name.split('.')
        
        if names[1] == 'txt':
        #f=#myfile.readlines()
            f = nltk.tokenize.sent_tokenize(myfile.read().decode())
            #f=re.findall( r'\w+|[^\s\w]+', )
            #print(f)
            #count=0
            #for i in lang:
            #  count+=1
            #print(count)
            #f=myfile.readlines()
            #save_path=''
            
            #translation = translator.translate("*] ^usg:", src='hi', dest='en')
            #print(translation.text)
            with open("media/files/"+myfile.name, 'w',encoding="utf-8") as op: 
                det=translator.detect(f[0:3])
                print(det)
                for line in f:
                    translation = translator.translate(line, src=src1, dest=dest1).text
                    if(det[0].lang==src1):   
                        for line in f:
                            translation = translator.translate(line, src=src1, dest=dest1).text
                            op.write(translation+' ')
                    #print(line+'\n')
                    else:
                        for line in f:
                            translation = translator.translate(line,dest=dest1).text
                            op.write(translation+' ')
            print(myfile.name)
            #sleep(30)
            #delete_file(myfile.name)
        #print ("Start : %s" % time.ctime())
        elif names[1]=='csv':
            print ("start : %s" % time.ctime())
            f=myfile.readlines()
            
            with open("media/files/"+myfile.name, 'w',encoding="utf-8") as op:      
                for line in f:
                        translation = translator.translate(line.decode(), src=src1, dest=dest1).text
                        op.write(translation)
                        op.write('\n')
            x=names[0]+".csv"
            #df=pd.read_table("static/files/"+names[0]+".txt", delimiter=',')
            #df.to_excel("static/files/"+names[0]+".xlsx")
            #filedf=pd.read_csv(myfile)
            #print(filedf)
            #workbook[i] = workbook[i].apply(translator.translate, src='en', dest='hi').apply(getattr, args=('text',))
            '''
            df = translator.translate(filedf,src='en',dest='hi')
            df.to_csv('my_file1.csv')
            '''
            #print ("End : %s" % time.ctime())
            #sleep(30)
            #delete_file(myfile.name)
        elif names[1]=='xlsx':
            fs = FileSystemStorage()
            upfile= fs.save("files/"+myfile.name, myfile)
            
            read_file = pd.read_excel("media/files/"+myfile.name)
            read_file.to_csv("media/files/"+names[0]+".csv",index = None,header=True)
           
            #delete_file(myfile.name)

            myfile=open("media/files/"+names[0]+".csv")
            f=myfile.readlines()

            translator=Translator()  
            with open("media/files/"+names[0]+".csv", 'w',encoding="utf-8") as op:      

                translation = translator.translate(f, src=src1,dest=dest1)
                for trans in translation:

                    op.write(trans.text)
                    op.write('\n')

            df_new=pd.read_csv("media/files/"+names[0]+".csv")
            GFG = pd.ExcelWriter("media/files/"+names[0]+"_translated."+names[1])
            df_new.to_excel(GFG, index=False)
            GFG.save() 
            x=names[0]+"_translated."+names[1]
            #delete_file(names[0]+'.csv')   
            #sleep(30)
            #delete_file(names[0]+"_translated."+names[1])
        
        elif names[1]=='pdf':

            #pdfFileObj = myfile.read() 
            print("Hello")
            #pdfReader = PyPDF2.PdfFileReader(io.BytesIO(pdfFileObj))
            #pdfwr = FPDF()
            #pages= pdfReader.numPages
            #docf = aw.Document('static/files/Cover Letter.pdf')
            #docf.save("Output.docx")
            #pages = [i for i in range(pages) pages.append(i)]
            #result = parse(pdf_file='static/files/Cover Letter.pdf', docx_with_path='static/files/'+names[0], pages=pages)
            '''
            for page in range(pdfReader.numPages):
                    pdfobj = pdfReader.getPage(page)
                    pdftext = pdfobj.extractText().split("\n")
                    pdfwr.add_page()
                    pdfwr.set_font("Arial", size = 10)
                    bullet_point = re.compile(r"\u2022")
                    for line in pdftext or ():
                        print(line)

                        translation = translator.translate(line,src=src1, dest=dest1).text
                        
                        translation = re.sub(bullet_point," ", translation)         
                        print(translation)

                        #line.replace(line, translation)

                        pdfwr.cell(100, 8 , txt = translation, ln = 1)
                        
                        
            pdfwr.output("media/files/"+ myfile.name)  
            '''
            '''
            fs = FileSystemStorage()
            upfile= fs.save("files/"+myfile.name, myfile)
            cvn = Converter("media/files/"+myfile.name)
            cvn.convert("media/files/"+names[0]+".docx")
            cvn.close()
            f=open("media/files/"+names[0]+".docx","rb")
            #f=myfile.read().decode()
            #f = docx.Document(myfile)
            b = open("media/files/"+names[0]+".html", 'w+', encoding="utf-8")

            custom_styles="""u => u
                             b => b
                             i => i
                            p[align='Center']=>p.center
                            table => table.table"""

            document = mammoth.convert_to_html(f,style_map=custom_styles)
            b.write(document.value)
            #print("beautiful soup 1")
            f.close()
            b.close()



            html = open("media/files/"+names[0]+".html",encoding="utf-8").read()
            soup = BeautifulSoup(html)
            tags = soup.find_all(["p","ul","ol","h1","h2","h3","h4","h5","h6","td","tr"])
            translator=Translator()
            
            for tag in tags:
            #constr="".join(str(e) for e in tag.contents)
            #print(constr)
                callcon(tag.contents, tag, src1,dest1)
            
            with open("media/files/"+names[0]+".html", "wb") as f_output:
                f_output.write(soup.prettify("utf-8"))



            try:
                # create the API client instance
                client = pdfcrowd.HtmlToPdfClient('DELTABOND007', '9c44a86338da91b0fe212625f209fe23')

                # run the conversion and write the result to a file
                client.convertFileToFile("media/files/"+names[0]+".html", "media/files/"+names[0]+"_translated.pdf")
                x=names[0]+"_translated."+names[1]
            except pdfcrowd.Error as why:
                # report the error
                sys.stderr.write('Pdfcrowd Error: {}\n'.format(why))

            '''
            '''
            fs = FileSystemStorage()
            upfile= fs.save("files/"+myfile.name, myfile)
            cvn = Converter("media/files/"+myfile.name)
            cvn.convert("media/files/"+names[0]+".docx")
            cvn.close() 
            f=open("media/files/"+names[0]+".docx","rb")
                        #f=myfile.read().decode()
                        #f = docx.Document(myfile)
            b = open("media/files/"+ names[0]+"html.html", 'w+', encoding="utf-8")

            custom_styles="""u => u
                                        b => b
                                        i => i
                                        p[align='Center']=>p.center
                                        table => table.table"""

            document = mammoth.convert_to_html(f,style_map=custom_styles)
            b.write(document.value)
                        #print("beautiful soup 1")
            f.close()
            b.close()


            htmlfile=open("media/files/"+ names[0]+"html.html", encoding="utf-8")
            html=htmlfile.read()
            soup = BeautifulSoup(html)
                        #print(soup)
            tags = soup.find_all(["p","ul","ol","h1","h2","h3","h4","h5","h6","td"])
                        #print(tags.get_text())
            for tag in tags:
                            #print(tag.get_text())
                        #constr="".join(str(e) for e in tag.contents)
                        #print(constr)
                callcon(tag.contents, tag, 'en', 'hi')
            htmlfile.close()
            with open("media/files/"+ names[0]+"html.html", "wb") as f_output:
                            f_output.write(soup.prettify("utf-8"))
                            #f_output.write("<p> </p>")

                    
                        
                        #os.environ.setdefault('PYPANDOC_PANDOC','/usr/bin/pandoc')

                        #pypandoc.download_pandoc()
                        #output = pypandoc.convert_file("media/files/"+names[0]+".html", format='html', to='docx', outputfile="media/files/"+names[0]+"_translated."+names[1], extra_args=['-RTS'])
            new_parser = HtmlToDocx()
            new_parser.table_style = 'TableGrid'
            new_parser.parse_html_file("media/files/"+ names[0]+"html.html","media/files/"+names[0]+"_translated")

            '''          


        elif names[1]=='docx':

           
            custom_css ="""
                        <style>
                        .red{
                            color: red;
                        }
                        .underline{
                            text-decoration: underline;
                        }
                        
                        table, th, td {
                        border: 1px solid black;
                        }
                        </style>"""
        
            fs = FileSystemStorage()
            upfile= fs.save("files/"+myfile.name, myfile)

            f=open("media/files/"+myfile.name,"rb")
            doc=docx.Document(f)
            det=translator.detect(doc.paragraphs[0].text)
            #print(doc.paragraphs[0])
            print(det.lang)
            if(det.lang==src1):
                src1=det.lang
            else:
                src1=det.lang
            #f=myfile.read().decode()
            #f = docx.Document(myfile)
            b = open("media/files/"+names[0]+".html", 'w+', encoding="utf-8")
            
            custom_styles="""u => u
                            b => b
                            i => i
                            p[align='Center']=>p.center
                            table => table.table"""

            document = mammoth.convert_to_html(f,style_map=custom_styles)
            b.write(document.value)
            #print("beautiful soup 1")
            f.close()
            b.close()
          
            #delete_file(myfile.name)
            htmlfile=open("media/files/"+names[0]+".html", encoding="utf-8")
            html=htmlfile.read()
            soup = BeautifulSoup(html)
            #print(soup)
            tags = soup.find_all(["p","ul","ol","h1","h2","h3","h4","h5","h6","td"])
            #print(tags.get_text())
                
            for tag in tags:
                callcon(tag.contents, tag, src1, dest1)
            htmlfile.close()
            with open("media/files/"+names[0]+".html", "wb") as f_output:
                f_output.write(soup.prettify("utf-8"))
                #f_output.write("<p> </p>")

           
            
            #os.environ.setdefault('PYPANDOC_PANDOC','/usr/bin/pandoc')

            #pypandoc.download_pandoc()
            #output = pypandoc.convert_file("media/files/"+names[0]+".html", format='html', to='docx', outputfile="media/files/"+names[0]+"_translated."+names[1], extra_args=['-RTS'])
            new_parser = HtmlToDocx()
            new_parser.table_style = 'TableGrid'
            new_parser.parse_html_file("media/files/"+names[0]+".html","media/files/"+names[0]+"_translated")
            x=names[0]+"_translated."+names[1]
            #pageObj = pdfReader.getPage(0)
            #delete_file(names[0]+".html")
            
            #extracting text from page
           
            
            
        
        #sleep(30)
        #delete_file(myfile.name)
        #timer = threading.Timer(600000, delete_file(myfile.name))
        #timer.start()
        
        #print ("End : %s" % time.ctime())
        #for line in f: 
        #print(myfile)
        #fs = FileSystemStorage()
        #filename = fs.save(myfile.name, myfile)
        truecheck = True
    return render(request,'home.html', {'lang': lang, 'truecheck': truecheck }) 

# def error_404(request, exception):
#     return render(request, '404.html')


# def error_404(request, exception):
#     return render(request, '404.html')
'''
def delete_file(myfile):
    if os.path.exists("media/files/"+myfile):
        os.remove("media/files/"+ myfile)
        #print("hello")
    else:
        print("The file does not exist")
'''
def callcon(t, tag, src1, dest1):
            for i in range(0, len(t)):
                if type(t[i])== NavigableString:
                    #print("Hello")
                    translator = Translator()
                    #print(t[i])
                    try:
                        translation = translator.translate(t[i], src=src1, dest=dest1).text
                    except:
                        translation=""
                    found = tag.find(text=str(t[i]))
                    if found:
                        new_text = found.replace_with(translation)
                    '''try:
                        new_text = tag.find(text=str(t[i])).replace_with(translation)
                    except:
                        
                        print(None)'''
                    #print(new_text)
                #elif t[i].name==img:

                else:
                    #print("Hello")
                    callcon(t[i].contents,tag, src1, dest1)

    



def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        pass

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def downloadfile(request):
    global x
    if x=="temp":
        messages.add_message(request, messages.INFO, 'File not uploaded. Please upload file')
        return HttpResponseRedirect('/')
    else:
        #print(x,"helloo")
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        filepath = base_dir + '/media/files/' + x
        thefile = filepath
        filename = os.path.basename(thefile)
        chunk_size = 8192
        response = StreamingHttpResponse(FileWrapper(open(thefile,'rb'),chunk_size),content_type=mimetypes.guess_type(thefile)[0])
        response['Content-Length'] = os.path.getsize(thefile)
        response['Content-Disposition'] = "Attachment;filename=%s" % filename
        
        #x = "temp"
        return response



def audiohome(request):
    '''
    global x
    #print(googletrans.LANGUAGES)
    lang =googletrans.LANGUAGES
    translator = Translator()
    if request.method == 'POST':
        myfile = request.FILES['myfile']
        x=myfile.name
        src1 = request.POST['src']
        dest1 = request.POST['dest']
        #intype=request.POST
        names = myfile.name.split('.')
        
        if names[1] == 'txt':
            
            file = myfile.read().decode()
            translation = translator.translate(file, src=src1, dest=dest1).text
            myobj=gTTS(text=translation,lang=dest1,slow=False)
            myobj.save("media/files/"+names[0]+".mp3")
        
        elif names[1]=='docx':
            doc = docx.Document(myfile)  # Creating word reader object.
            data = ""
            fullText = []
            translator= Translator()
            try:
                    for para in doc.paragraphs: 
                        translation=translator.translate(para.text, src=src1, dest=dest1).text
                        fullText.append(translation)
                        data = '\n'.join(fullText)
            
            except:
                    print('There was an error opening the file!')
                    

            #language='en'

            myobj=gTTS(text=data,lang=dest1,slow=False)
            myobj.save("media/files/"+names[0]+".mp3")
    '''
    return render(request,'home.html')
            #os.system("welcome.mp3")
